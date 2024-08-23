from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches
import os

app = Flask(__name__)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Collecting form data
        referring_doctor = request.form.get("referring_doctor")
        patient_name = request.form.get("patient_name")
        nhs_number = request.form.get("nhs_number")
        date_of_clinic = request.form.get("date_of_clinic")
        store_name = request.form.get("store_name")
        store_phone_number = request.form.get("store_phone_number")
        audiologist = request.form.get("audiologist")
        history = request.form.get("history")
        otoscopy = request.form.get("otoscopy")
        pta = request.form.get("pta")
        individual_management_plan = request.form.get("individual_management_plan")
        additional_notes = request.form.get("additional_notes")

        # Create a Word Document
        doc = Document()

        # Add centered logo
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = 1  # Center alignment
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture('images/Specsaversaudio-removebg-preview.png', width=Inches(2.5))

        # Store details beneath the logo
        store_info_paragraph = doc.add_paragraph()
        store_info_paragraph.alignment = 1  # Center alignment
        store_info_paragraph.add_run(f"{store_name}\n{store_phone_number}\n").bold = True

        doc.add_paragraph("\n")  # Add a blank line

        # Patient details section
        doc.add_heading('Patient Details', level=2)
        patient_details = doc.add_paragraph()
        patient_details.add_run(f"Patient Name: {patient_name}\n").bold = True
        patient_details.add_run(f"NHS Number: {nhs_number}\n").bold = True
        patient_details.add_run(f"Referring GP: {referring_doctor}\n").bold = True
        patient_details.add_run(f"Clinic Date: {date_of_clinic}\n").bold = True

        doc.add_paragraph("\n")  # Add a blank line

        # Introduction text
        doc.add_paragraph(
            "Thank you for referring the above patient to Specsavers Hearcare. "
            "A full audiological assessment was undertaken and the results were as follows:\n"
        )

        # Clinical details
        doc.add_heading('History', level=2)
        doc.add_paragraph(history)

        doc.add_heading('Otoscopy', level=2)
        doc.add_paragraph(otoscopy)

        doc.add_heading('PTA', level=2)
        doc.add_paragraph(pta)

        doc.add_heading('Individual Management Plan', level=2)
        doc.add_paragraph(individual_management_plan)

        doc.add_paragraph("\n")  # Add a blank line

        # Closing text
        doc.add_paragraph(
            f"Please contact the service on {store_phone_number} if you need any further information."
        )

        doc.add_paragraph("\n")  # Add a blank line

        # Signature
        doc.add_paragraph(f"Yours sincerely,\n\n{audiologist}\nAudiologist")

        # Save the document
        output_filename = "clinic_letter.docx"
        doc.save(output_filename)

        # Send the file to the user
        return send_file(output_filename, as_attachment=True)

    return render_template("index.html")


if __name__ == "__main__":
    # Ensure the images folder exists
    if not os.path.exists('images'):
        os.makedirs('images')
    app.run(debug=True)

